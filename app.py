import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import datetime

st.set_page_config(page_title="Excel → Word (mẫu)", layout="wide")

st.title("Chuyển dữ liệu Excel thành file Word theo mẫu")

uploaded_file = st.file_uploader("Tải lên file Excel (.xlsx) chứa danh sách", type=["xlsx","xls"], accept_multiple_files=False)

def detect_columns(df):
    # mapping candidates: key = normalized name, value = actual column in df
    cols = {c.lower().strip(): c for c in df.columns}
    def find(keys):
        for k in keys:
            if k.lower() in cols:
                return cols[k.lower()]
        # try partial match
        for c in df.columns:
            low = c.lower()
            for k in keys:
                if k.lower() in low:
                    return c
        return None
    mapping = {
        "ho_ten": find(["họ và tên","họ và tên","họ tên","ho va ten","ho ten","họ và tên ngà","họ và tên ngày"]),
        "ngay_sinh": find(["ngày tháng năm sinh","ngay thang nam sinh","ngày sinh","ngay sinh"]),
        "don_vi": find(["đv","đơn vị","don vi","đơn vi"]),
        "cb_cv": find(["cb","cb cv","cb cv","cb cv","cb cv"]),
        "nhap_ngu": find(["n.n","n.n.","nhập ngũ","nhap ngũ","nhập ngu"]),
        "dan_toc": find(["dân tộc","dân toc","dan toc"]),
        "van_hoa": find(["văn hóa","văn hoa","van hoa"]),
        "so_cccd": find(["số cccd","số cmt","số cmnd","so cccd","cccd"]),
        "que_quan": find(["quê quán","quê quán trú quán","quê quán trú","quê quán trú quán","quê quán/trú quán","quê quán/trú quán"]),
        "xa": find(["xã","xa"]),
        "tinh": find(["tỉnh","tinh"]),
        "bo": find(["bố","bo"]),
        "me": find(["mẹ","me"]),
        "sdt_gia_dinh": find(["sdt","số điện thoại gia đình","sd t gia đình","sd t gia dinh","sdt gia đình","sdt gia dinh"])
    }
    return mapping

def format_cell(val):
    if pd.isna(val):
        return ""
    if isinstance(val, float) and val.is_integer():
        return str(int(val))
    return str(val)

def create_word_from_df(df):
    mapping = detect_columns(df)
    # ensure columns exist; if not, create empty
    for v in mapping.values():
        if v is None:
            # nothing to do; we'll use empty values
            pass

    # Normalize columns into a working DataFrame
    work = pd.DataFrame()
    work["Họ và tên"] = df[mapping["ho_ten"]] if mapping["ho_ten"] in df.columns else ""
    work["Ngày sinh"] = df[mapping["ngay_sinh"]] if mapping["ngay_sinh"] in df.columns else ""
    work["Đơn vị"] = df[mapping["don_vi"]] if mapping["don_vi"] in df.columns else ""
    work["Cb CV"] = df[mapping["cb_cv"]] if mapping["cb_cv"] in df.columns else ""
    work["Nhập ngũ"] = df[mapping["nhap_ngu"]] if mapping["nhap_ngu"] in df.columns else ""
    work["Dân tộc"] = df[mapping["dan_toc"]] if mapping["dan_toc"] in df.columns else ""
    work["Văn hóa"] = df[mapping["van_hoa"]] if mapping["van_hoa"] in df.columns else ""
    work["Số CCCD"] = df[mapping["so_cccd"]] if mapping["so_cccd"] in df.columns else ""
    work["Quê quán"] = df[mapping["que_quan"]] if mapping["que_quan"] in df.columns else ""
    work["Xã"] = df[mapping["xa"]] if mapping["xa"] in df.columns else ""
    work["Tỉnh"] = df[mapping["tinh"]] if mapping["tinh"] in df.columns else ""
    work["Bố"] = df[mapping["bo"]] if mapping["bo"] in df.columns else ""
    work["Mẹ"] = df[mapping["me"]] if mapping["me"] in df.columns else ""
    work["SDT gia đình"] = df[mapping["sdt_gia_dinh"]] if mapping["sdt_gia_dinh"] in df.columns else ""

    # Fill NaN with empty string
    work = work.fillna("")

    # Group by Tỉnh then Xã
    grouped = work.groupby(["Tỉnh","Xã"], sort=False)

    doc = Document()
    # set default font (Times New Roman) and size
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(12)

    # Header (you can customize)
    doc_para = doc.add_paragraph()
    run = doc_para.add_run("TỔNG HỢP TRÍCH NGANG")
    run.bold = True
    run.font.size = Pt(14)
    doc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    current_province = None
    for (province, xa), group in grouped:
        province = format_cell(province)
        xa = format_cell(xa)
        if province and province != current_province:
            doc.add_paragraph()  # spacing
            p = doc.add_paragraph(province)
            p.runs[0].bold = True
            current_province = province
        if xa:
            p2 = doc.add_paragraph(f"{xa}")
            p2.runs[0].italic = True

        # create table with header row matching your Word sample
        table = doc.add_table(rows=1, cols=15)
        hdr_cells = table.rows[0].cells
        headers = ["TT","Họ và tên Ngày tháng năm sinh","Cb CV","Đơn vị","Nhập ngũ","Thành phần","Văn hóa","Sức khỏe","DT TG","Ngày vào Đoàn","Ngày vào Đảng","Quê quán Trú Quán","Họ tên bố, mẹ Họ tên vợ, con SỐ ĐIỆN THOẠI GIA ĐÌNH","Số cccd","Ghi chú"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        # fill rows
        for idx, row in group.reset_index(drop=True).iterrows():
            r = table.add_row().cells
            r[0].text = str(idx+1)
            r[1].text = f"{format_cell(row['Họ và tên'])} {format_cell(row['Ngày sinh'])}"
            r[2].text = format_cell(row["Cb CV"])
            r[3].text = format_cell(row["Đơn vị"])
            r[4].text = format_cell(row["Nhập ngũ"])
            # placeholders for columns not in excel; leave empty or adapt
            r[5].text = ""  # Thành phần
            r[6].text = format_cell(row["Văn hóa"])
            r[7].text = ""  # Sức khỏe
            r[8].text = format_cell(row["Dân tộc"])
            r[9].text = ""  # Ngày vào Đoàn
            r[10].text = ""  # Ngày vào Đảng
            r[11].text = format_cell(row["Quê quán"])
            r[12].text = f"{format_cell(row['Bố'])} {format_cell(row['Mẹ'])} {format_cell(row['SDT gia đình'])}"
            r[13].text = format_cell(row["Số CCCD"])
            r[14].text = ""  # Ghi chú

    # Save to BytesIO and return
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

if uploaded_file is not None:
    try:
        # read all sheets and concat
        xls = pd.ExcelFile(uploaded_file)
        df_list = []
        for sh in xls.sheet_names:
            tmp = pd.read_excel(xls, sheet_name=sh, dtype=str)
            if tmp.shape[1] > 0:
                df_list.append(tmp)
        if not df_list:
            st.error("Không tìm thấy dữ liệu trong file Excel.")
        else:
            df = pd.concat(df_list, ignore_index=True)
            st.success(f"Đã đọc Excel: {df.shape[0]} dòng, {df.shape[1]} cột.")
            st.dataframe(df.head(10))

            if st.button("Tạo file Word từ dữ liệu"):
                with st.spinner("Đang tạo file Word..."):
                    doc_bytes = create_word_from_df(df)
                    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"tong_hop_trich_ngang_{now}.docx"
                    st.download_button("Tải file Word", data=doc_bytes, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(f"Lỗi khi đọc file Excel: {e}")
else:
    st.info("Vui lòng tải file Excel lên để bắt đầu.")
